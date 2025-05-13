import sqlite3
import os
import random
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    Application,
    CommandHandler,
    ConversationHandler,
    MessageHandler,
    CallbackQueryHandler,
    filters,
    ContextTypes,
)
from docx import Document
import re
from datetime import datetime

# مسیر فایل‌ها
def resource_path(relative_path):
    base_path = os.path.dirname(os.path.abspath(__file__))
    print(f"Base path: {base_path}")  # دیباگ مسیر پایه
    return os.path.join(base_path, relative_path)

# خواندن فایل ورد
def read_docx(file_path):
    file_path = resource_path(file_path)
    print(f"Attempting to read file: {file_path}")  # دیباگ
    if not os.path.exists(file_path):
        print(f"Error: File {file_path} not found!")  # دیباگ
        return f"خطا: فایل {file_path} یافت نشد!"
    try:
        doc = Document(file_path)
        full_text = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        print(f"Successfully read file: {file_path}")  # دیباگ
        return "\n".join(full_text)
    except Exception as e:
        print(f"Error reading docx file {file_path}: {e}")  # دیباگ
        return f"خطا در خواندن فایل ورد: {e}"

# خواندن input.docx برای دیتابیس
def read_docx_for_db(file_path):
    file_path = resource_path(file_path)
    print(f"Attempting to read file for DB: {file_path}")  # دیباگ
    if not os.path.exists(file_path):
        print(f"Error: File {file_path} not found!")  # دیباگ
        return []
    try:
        with open(file_path, "rb") as f:
            doc = Document(f)
        full_text = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        print(f"Successfully read file for DB: {file_path}")  # دیباگ
        return full_text
    except Exception as e:
        print(f"Error reading docx file for DB {file_path}: {e}")  # دیباگ
        return []

# جدا کردن بخش‌ها
def split_sections(lines):
    sections = []
    current_section = []
    for line in lines:
        if line.startswith("شماره صفحه:"):
            if current_section:
                sections.append(current_section)
            current_section = [line]
        else:
            current_section.append(line)
    if current_section:
        sections.append(current_section)
    return sections

# پارس کردن بخش
def parse_section(section):
    record = {
        "page_number": "",
        "surah_name": "",
        "verse_number": "",
        "verse_text": "",
        "translation": "",
        "istikhara_result": ""
    }
    try:
        for line in section:
            if line.startswith("شماره صفحه:"):
                match = re.search(r"\d+", line)
                record["page_number"] = match.group() if match else ""
            elif line.startswith("نام سوره:"):
                record["surah_name"] = line.split(":", 1)[1].strip()
            elif line.startswith("شماره آیه:"):
                match = re.search(r"\d+", line)
                record["verse_number"] = match.group() if match else ""
            elif line.startswith("متن آیه:"):
                record["verse_text"] = line.split(":", 1)[1].strip()
            elif line.startswith("ترجمه (فولادوند):"):
                record["translation"] = line.split(":", 1)[1].strip()
            elif line.startswith("نتیجه استخاره شما:"):
                record["istikhara_result"] = line.split(":", 1)[1].strip()
            else:
                record["verse_text"] += " " + line.strip() if record["verse_text"] else line.strip()
    except Exception as e:
        print(f"Error parsing section: {e}")  # دیباگ
    return record

# ایجاد یا به‌روزرسانی دیتابیس
def create_database():
    data_dir = resource_path("data")
    db_path = os.path.join(data_dir, "quran_istikhara.db")
    print(f"Database path: {db_path}")  # دیباگ
    
    if not os.path.exists(data_dir):
        print(f"Creating data directory: {data_dir}")  # دیباگ
        os.makedirs(data_dir)

    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS istikhara (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            page_number TEXT,
            surah_name TEXT,
            verse_number TEXT,
            verse_text TEXT,
            translation TEXT,
            istikhara_result TEXT
        )
    ''')

    cursor.execute("SELECT COUNT(*) FROM istikhara")
    if cursor.fetchone()[0] == 0:
        file_path = resource_path("data/input.docx")
        lines = read_docx_for_db(file_path)
        if not lines:
            print("No data read from input.docx!")  # دیباگ
            conn.close()
            return
        sections = split_sections(lines)
        records = [parse_section(section) for section in sections]
        for record in records:
            cursor.execute('''
                INSERT INTO istikhara (page_number, surah_name, verse_number, verse_text, translation, istikhara_result)
                VALUES (?, ?, ?, ?, ?, ?)
            ''', (
                record["page_number"],
                record["surah_name"],
                record["verse_number"],
                record["verse_text"],
                record["translation"],
                record["istikhara_result"]
            ))
        print(f"Database successfully populated from {file_path}")  # دیباگ

    conn.commit()
    conn.close()

# ثابت‌ها برای Conversation Handler
INTENT, SEARCH_PAGE = range(2)

# منوی اصلی
def get_main_menu():
    keyboard = [
        [InlineKeyboardButton("ادامه", callback_data="continue")],
        [InlineKeyboardButton("جستجو", callback_data="search")],
        [InlineKeyboardButton("تاریخچه", callback_data="history")]
    ]
    return InlineKeyboardMarkup(keyboard)

# تابع شروع
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    print("Received /start command")  # دیباگ
    create_database()
    description = read_docx("data/discription.docx")
    for chunk in [description[i:i+4096] for i in range(0, len(description), 4096)]:
        await update.message.reply_text(
            chunk,
            reply_markup=get_main_menu() if chunk == description[-4096:] else None
        )

# تابع ادامه
async def continue_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    edameh = read_docx("data/edameh.docx")
    keyboard = [[InlineKeyboardButton("نیت", callback_data="intent")]]
    reply_markup = InlineKeyboardMarkup(keyboard)
    for chunk in [edameh[i:i+4096] for i in range(0, len(edameh), 4096)]:
        await query.message.reply_text(
            chunk,
            reply_markup=reply_markup if chunk == edameh[-4096:] else None
        )

# تابع نیت
async def intent_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    niyat = read_docx("data/niyat.docx")
    await query.message.reply_text(
        f"{niyat}\n\nلطفاً نیت خود را بنویسید:",
        reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("لغو", callback_data="cancel")]])
    )
    return INTENT

# دریافت نیت
async def receive_intent(update: Update, context: ContextTypes.DEFAULT_TYPE):
    intent = update.message.text.strip()
    if not intent:
        await update.message.reply_text("لطفاً نیت خود را بنویسید:")
        return INTENT

    context.user_data["intent"] = intent
    await update.message.reply_text("در حال انجام استخاره... لطفاً منتظر بمانید")
    
    db_path = resource_path("data/quran_istikhara.db")
    print(f"Accessing database: {db_path}")  # دیباگ
    try:
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM istikhara")
        rows = cursor.fetchall()
        if not rows:
            await update.message.reply_text("خطا: دیتابیس خالی است!")
            conn.close()
            return ConversationHandler.END
        
        record = random.choice(rows)
        conn.close()
    except Exception as e:
        print(f"Error accessing database {db_path}: {e}")  # دیباگ
        await update.message.reply_text(f"خطا در دسترسی به دیتابیس: {e}")
        return ConversationHandler.END

    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    result_text = (
        f"#استخاره\n"
        f"نیت: {intent}\n"
        f"زمان: {timestamp}\n"
        f"شماره صفحه: {record[1]}\n"
        f"نام سوره: {record[2]}\n"
        f"شماره آیه: {record[3]}\n"
        f"متن آیه:\n{record[4]}\n"
        f"ترجمه:\n{record[5]}\n"
        f"نتیجه استخاره: {record[6]}"
    )

    keyboard = [[InlineKeyboardButton("بازگشت به منوی اصلی", callback_data="main_menu")]]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text(result_text, reply_markup=reply_markup)

    return ConversationHandler.END

# تابع جستجو
async def search_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    await query.message.reply_text(
        "لطفاً شماره صفحه‌ای که می‌خواهید را بنویسید:",
        reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("لغو", callback_data="cancel")]])
    )
    return SEARCH_PAGE

# دریافت شماره صفحه
async def receive_page(update: Update, context: ContextTypes.DEFAULT_TYPE):
    page_number = update.message.text.strip()
    if not page_number.isdigit():
        await update.message.reply_text("لطفاً یک شماره صفحه معتبر بنویسید:")
        return SEARCH_PAGE

    db_path = resource_path("data/quran_istikhara.db")
    print(f"Accessing database for search: {db_path}")  # دیباگ
    try:
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM istikhara WHERE page_number = ?", (page_number,))
        rows = cursor.fetchall()
        conn.close()
    except Exception as e:
        print(f"Error accessing database for search {db_path}: {e}")  # دیباگ
        await update.message.reply_text(f"خطا در دسترسی به دیتابیس: {e}")
        return ConversationHandler.END

    if not rows:
        await update.message.reply_text(
            f"هیچ رکوردی برای شماره صفحه {page_number} یافت نشد!",
            reply_markup=get_main_menu()
        )
        return ConversationHandler.END

    result_text = ""
    for row in rows:
        result_text += (
            f"شماره صفحه: {row[1]}\n"
            f"نام سوره: {row[2]}\n"
            f"شماره آیه: {row[3]}\n"
            f"متن آیه:\n{row[4]}\n"
            f"ترجمه:\n{row[5]}\n"
            f"نتیجه استخاره: {row[6]}\n"
            f"{'-' * 50}\n"
        )

    keyboard = [[InlineKeyboardButton("بازگشت به منوی اصلی", callback_data="main_menu")]]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text(result_text, reply_markup=reply_markup)
    return ConversationHandler.END

# نمایش تاریخچه
async def history_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    await query.message.reply_text(
        "تاریخچه استخاره‌های شما در پیام‌های قبلی با تگ #استخاره ذخیره شده‌اند. لطفاً پیام‌های چت را بررسی کنید.",
        reply_markup=get_main_menu()
    )

# لغو مکالمه
async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    await query.message.reply_text("عملیات لغو شد.", reply_markup=get_main_menu())
    return ConversationHandler.END

# بازگشت به منوی اصلی
async def main_menu(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    await query.message.reply_text("به منوی اصلی خوش آمدید:", reply_markup=get_main_menu())

# خطا
async def error_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    print(f"Update {update} caused error {context.error}")  # دیباگ
    if update.message:
        await update.message.reply_text("خطایی رخ داد. لطفاً دوباره امتحان کنید.", reply_markup=get_main_menu())

def main():
    print("Starting bot...")  # دیباگ
    TOKEN = os.getenv("TOKEN")
    print(f"Token: {'Set' if TOKEN else 'Not set'}")  # دیباگ
    if not TOKEN:
        raise ValueError("توکن ربات در متغیر محیطی TOKEN پیدا نشد!")
    
    try:
        application = Application.builder().token(TOKEN).build()
        print("Application initialized successfully")  # دیباگ
    except Exception as e:
        print(f"Error initializing application: {e}")  # دیباگ
        raise

    intent_conv = ConversationHandler(
        entry_points=[CallbackQueryHandler(intent_handler, pattern="intent")],
        states={
            INTENT: [MessageHandler(filters.TEXT & ~filters.COMMAND, receive_intent)],
        },
        fallbacks=[CallbackQueryHandler(cancel, pattern="cancel")]
    )

    search_conv = ConversationHandler(
        entry_points=[CallbackQueryHandler(search_handler, pattern="search")],
        states={
            SEARCH_PAGE: [MessageHandler(filters.TEXT & ~filters.COMMAND, receive_page)],
        },
        fallbacks=[CallbackQueryHandler(cancel, pattern="cancel")]
    )

    application.add_handler(CommandHandler("start", start))
    application.add_handler(intent_conv)
    application.add_handler(search_conv)
    application.add_handler(CallbackQueryHandler(continue_handler, pattern="continue"))
    application.add_handler(CallbackQueryHandler(history_handler, pattern="history"))
    application.add_handler(CallbackQueryHandler(main_menu, pattern="main_menu"))
    application.add_error_handler(error_handler)

    print("Bot handlers set up, starting polling...")  # دیباگ
    application.run_polling()

if __name__ == "__main__":
    main()